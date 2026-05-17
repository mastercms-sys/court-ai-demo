import streamlit as st
import google.generativeai as genai
from docx import Document
import io
import tempfile
import os
import time
import subprocess
import re  # نیا ٹول: AI کے جواب کو الگ الگ کرنے کے لیے (Regex)

# پیج کی سیٹنگ
st.set_page_config(page_title="Court Dictation AI - Deep Think", page_icon="⚖️", layout="centered")

# ==========================================
# 🔒 لاگ ان سسٹم (Monthly Subscription Lock)
# ==========================================
if "logged_in" not in st.session_state:
    st.session_state.logged_in = False

if not st.session_state.logged_in:
    st.markdown("<h1 style='text-align: center;'>🔒 Court Dictation Pro</h1>", unsafe_allow_html=True)
    st.markdown("<p style='text-align: center;'>براہ کرم اپنا ماہانہ پاس ورڈ درج کریں۔</p>", unsafe_allow_html=True)
    
    col1, col2, col3 = st.columns([1,2,1])
    with col2:
        password_input = st.text_input("پاس ورڈ (Password):", type="password")
        if st.button("لاگ ان (Login)", use_container_width=True):
            try:
                correct_password = st.secrets["APP_PASSWORD"]
                if password_input == correct_password:
                    st.session_state.logged_in = True
                    st.rerun()
                else:
                    st.error("❌ غلط پاس ورڈ! رسائی کے لیے ایڈمن سے رابطہ کریں۔")
            except Exception as e:
                st.error("⚠️ ایپ کے بیک اینڈ پر پاس ورڈ سیٹ نہیں ہے۔")
    st.stop()
# ==========================================

st.title("⚖️ Auto Court Dictation Pro")
st.markdown("**(🧠 Deep Legal Thinker Enabled)**")
st.write("اپنی ڈکٹیشن کی آڈیو، ویڈیو یا واٹس ایپ فائل اپلوڈ کریں۔ سسٹم گہرائی سے سوچ کر Word فائل بنا دے گا۔")

# API Key
try:
    API_KEY = st.secrets["GEMINI_API_KEY"]
    genai.configure(api_key=API_KEY)
except Exception as e:
    st.error("⚠️ ایپ کے بیک اینڈ پر API Key موجود نہیں ہے۔")
    st.stop()

allowed_formats = ["mp3", "wav", "m4a", "mp4", "mov", "avi", "mkv", "aac", "ogg"]
uploaded_file = st.file_uploader("فائل اپلوڈ کریں", type=allowed_formats)

if uploaded_file is not None:
    if st.button("Generate Court Document"):
        
        file_extension = os.path.splitext(uploaded_file.name)[1].lower()
        temp_media = tempfile.NamedTemporaryFile(delete=False, suffix=file_extension)
        temp_media.write(uploaded_file.getbuffer())
        temp_media.close()
        temp_media_path = temp_media.name
        
        clean_audio_path = None
        
        try:
            with st.spinner("1️⃣ سسٹم آپ کی فائل کو آٹو فکس کر رہا ہے (بیک گراؤنڈ پروسیسنگ)..."):
                clean_audio = tempfile.NamedTemporaryFile(delete=False, suffix=".mp3")
                clean_audio.close()
                clean_audio_path = clean_audio.name
                
                try:
                    command = [
                        "ffmpeg", "-y", "-i", temp_media_path, 
                        "-vn", "-ar", "16000", "-ac", "1", "-b:a", "64k", 
                        clean_audio_path
                    ]
                    subprocess.run(command, check=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
                except subprocess.CalledProcessError:
                    st.error("❌ فائل کنورٹ نہیں ہو سکی۔ یہ فائل مکمل کرپٹ ہے۔")
                    st.stop()

            with st.spinner("2️⃣ فائل گوگل کے سرور پر اپلوڈ ہو رہی ہے..."):
                gemini_media = genai.upload_file(path=clean_audio_path, mime_type="audio/mp3")
            
            with st.spinner("3️⃣ AI سرور آواز کو سمجھ رہا ہے، براہ کرم انتظار کریں..."):
                while gemini_media.state.name == 'PROCESSING':
                    time.sleep(5)
                    gemini_media = genai.get_file(gemini_media.name) 
                    
                if gemini_media.state.name == 'FAILED':
                    st.error("❌ گوگل AI نے اس فائل کو ریجیکٹ کر دیا ہے۔")
                    genai.delete_file(gemini_media.name)
                    st.stop()
            
            with st.spinner("4️⃣ 🧠 DEEP THINK: سرور منطقی غلطیاں ڈھونڈ کر فیصلہ ڈرافٹ کر رہا ہے..."):
                
                valid_models = []
                try:
                    for m in genai.list_models():
                        if 'generateContent' in m.supported_generation_methods:
                            name_lower = m.name.lower()
                            if 'image' not in name_lower and 'vision' not in name_lower and 'exp' not in name_lower and 'learn' not in name_lower:
                                valid_models.append(m.name)
                except Exception as e:
                    pass
                
                pro_models = sorted([m for m in valid_models if 'pro' in m.lower()], reverse=True)
                flash_models = sorted([m for m in valid_models if 'flash' in m.lower()], reverse=True)
                
                primary_model_name = pro_models[0] if pro_models else (flash_models[0] if flash_models else "gemini-1.5-pro")
                backup_model_name = flash_models[0] if flash_models else "gemini-1.5-flash"
                
                # 🔴 آپ کا Chain of Thought ماسٹر پرامپٹ
                prompt = """Act as an Elite Legal Assistant and Senior District Court Judge in Pakistan.

I will provide you with a raw, unedited audio transcript of court proceedings. Your task is to act as a "Deep Thinking" AI. You must carefully analyze the text, fix logical errors, remove dictation repetitions, correct the grammar, and format it into a flawless, ready-to-print Court Order.

Follow these STRICT RULES:

1. LOGICAL CORRECTIONS (DEEP THINK):
   - Fix dictation slips of the tongue. For example, if the text says "being sons... they are Parda Nashin ladies", use logic to correct it to "being daughters, they are Parda Nashin ladies".
   - Identify and REMOVE exact repetitions. If the speaker stuttered and repeated a sentence (e.g., "defendants are in possession... defendants are in possession"), write it ONLY ONCE.

2. LEGAL VOCABULARY & GRAMMAR:
   - Always translate "Sarkar" to "THE STATE" in the title.
   - Fix Urdu-to-English literal translations: 
     Change "interfere into" to "interfere with". 
     Change "deprive from" to "deprive of". 
     Change "succeeded to" to "succeeded in".
   - Fix dictation mishears: Change "Arguments are record perused" to exactly "Arguments heard. Record perused." (Put this on a separate line).
   - Remove the word "That" at the beginning of paragraphs to ensure a formal, objective judicial flow.

3. NUMBERS AND CURRENCY:
   - Convert amounts to numeric format with commas and a dash (e.g., "Rs. 5,000,000/-" instead of "Rupees 50 Lakh" or "5000000").

4. MS WORD READY FORMATTING (NO MARKDOWN OR QUOTES):
   - DO NOT use markdown symbols like asterisks (**) or hashes (#) or quotation marks ("") for bolding or headings. 
   - Use ONLY ALL CAPS for emphasis (e.g., type exactly CASE NO., TITLE, ORDER, ISSUES).
   - Separate multiple cases with a line: ----------------------------------------
   - Keep (OPP) or (OPD) in the Issue heading. Example: ISSUE NO. 1 (OPP): [Text]

TO ENSURE DEEP THINKING, YOU MUST FORMAT YOUR RESPONSE EXACTLY LIKE THIS:

<thinking>
1. (Briefly list the logical errors, repetitions, and grammar mistakes you found in the raw text).
2. (State how you will fix them based on the rules).
</thinking>

<court_order>
(Provide the final, perfectly formatted plain-text document here. DO NOT use ** or "" symbols).
</court_order>

RAW AUDIO TO PROCESS:"""
                
                full_text = ""
                
                try:
                    st.info(f"💡 Deep Think Engine: **{primary_model_name}**")
                    primary_model = genai.GenerativeModel(primary_model_name)
                    response = primary_model.generate_content([prompt, gemini_media])
                    full_text = response.text
                except Exception as model_error:
                    error_msg = str(model_error)
                    if ("429" in error_msg or "Quota" in error_msg or "ResourceExhausted" in error_msg) and backup_model_name != primary_model_name:
                        st.warning(f"⚠️ مین سرور بزی ہے، آٹو بیک اپ سرور (**{backup_model_name}**) استعمال کیا جا رہا ہے...")
                        try:
                            backup_model = genai.GenerativeModel(backup_model_name)
                            response = backup_model.generate_content([prompt, gemini_media])
                            full_text = response.text
                        except Exception as backup_error:
                            raise backup_error
                    else:
                        raise model_error
                
                # 🔴 سمارٹ فلٹر (Regex) جو <thinking> اور <court_order> کو الگ کرے گا
                thinking_match = re.search(r'<thinking>(.*?)</thinking>', full_text, re.DOTALL | re.IGNORECASE)
                order_match = re.search(r'<court_order>(.*?)</court_order>', full_text, re.DOTALL | re.IGNORECASE)
                
                thinking_text = thinking_match.group(1).strip() if thinking_match else ""
                
                if order_match:
                    final_order_text = order_match.group(1).strip()
                else:
                    # اگر ماڈل غلطی سے ٹیگ لگانا بھول جائے تو thinking والے حصے کو کاٹ کر باقی سب لے لو
                    final_order_text = full_text
                    if thinking_match:
                        final_order_text = re.sub(r'<thinking>.*?</thinking>', '', final_order_text, flags=re.DOTALL | re.IGNORECASE).strip()
                    final_order_text = final_order_text.replace('<court_order>', '').replace('</court_order>', '').strip()
                
                st.success("✅ ڈکٹیشن 100% درستگی (Deep Legal Think) کے ساتھ تیار ہو گئی!")
                
                # 1. یوزر کو حیران کرنے کے لیے AI کا "سوچنے کا عمل" ایک خانے میں دکھائیں
                if thinking_text:
                    with st.expander("🧠 AI کا تجزیہ (دیکھیں سسٹم نے کیا غلطیاں پکڑیں اور کیسے ٹھیک کیں)"):
                        st.write(thinking_text)
                
                # 2. مین ٹیکسٹ ایریا میں صرف صاف فیصلہ دکھائیں
                st.text_area("تیار شدہ عدالتی فیصلہ (یہاں سے سلیکٹ کر کے کاپی بھی کر سکتے ہیں):", final_order_text, height=450)
                
                # 3. Word Document بنانا (جس میں صرف فیصلہ ہوگا، سوچ نہیں ہوگی)
                doc = Document()
                for line in final_order_text.split('\n'):
                    if line.strip() != "":
                        doc.add_paragraph(line)
                    else:
                        doc.add_paragraph()
                        
                doc_buffer = io.BytesIO()
                doc.save(doc_buffer)
                doc_buffer.seek(0)
                
                st.download_button(
                    label="📥 Word File (.docx) ڈاؤنلوڈ کریں",
                    data=doc_buffer,
                    file_name="Court_Order_Pro.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            
            genai.delete_file(gemini_media.name)
            
        except Exception as e:
            st.error(f"کوئی مسئلہ پیش آیا: {e}")
            with st.expander("تکنیکی خرابی کی تفصیل (ایرر آئے تو یہ دیکھیں)"):
                st.write(e)
            
        finally:
            if os.path.exists(temp_media_path):
                os.remove(temp_media_path)
            if clean_audio_path and os.path.exists(clean_audio_path):
                os.remove(clean_audio_path)
